import sys
import os
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from location_of_services import address


class GfeDocument:
    """Creates a Good Faith Estimate form and saves it in docx."""

    def __init__(
        self,
        section1: "text file",
        section2: "text file",
        estimate_info: "Object",
        session_count_low=12,
        session_count_high=24,
    ):
        self.section1 = self.resource_path(section1)
        self.section2 = self.resource_path(section2)
        self.estimate_info = estimate_info
        self.session_count_low = session_count_low
        self.session_count_high = session_count_high
        self.filename = None

        registration_fee = (
            "Registr-\nation fee",
            "None",
            "None",
            "$25",
            "1",
            "$25",
        )

        low_estimate = [
            (registration_fee),
            (
                "Initial evaluation",
                "90971",
                "None",
                f"${str(self.estimate_info.session_rate)}",
                "1",
                f"${str(self.estimate_info.session_rate)}",
            ),
            (
                "Psycho-\ntherapy",
                str(self.estimate_info.services_sought),
                "None",
                f"${str(self.estimate_info.session_rate)}",
                str(session_count_low),
                f"${str(self.calculate_low_estimate())}",
            ),
        ]

        high_estimate = [
            (registration_fee),
            (
                "Initial evaluation",
                "90971",
                "None",
                f"${str(self.estimate_info.session_rate)}",
                "1",
                f"${str(self.estimate_info.session_rate)}",
            ),
            (
                "Psycho-\ntherapy",
                str(self.estimate_info.services_sought),
                "None",
                f"${str(self.estimate_info.session_rate)}",
                str(self.session_count_high),
                f"${str(self.calculate_high_estimate())}",
            ),
        ]

        document = Document()

        heading = document.add_heading(
            "Life Resources Good Faith Estimate", level=1
        )
        heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(self.create_section1())
        document.add_page_break()
        if self.estimate_info.estimate_type == "Renewal":
            document.add_heading(
                "Itemized estimate for 12 session course of treatment", level=2
            )
            self.create_estimate_table(document, low_estimate[2:])
            document.add_heading(
                "Itemized estimate for 24 session course of treatment", level=2
            )
            self.create_estimate_table(document, high_estimate[2:])
            document.add_paragraph(
                f"\nEstimate range: ${self.calculate_low_estimate()}-"
                f"{self.calculate_high_estimate()}"
                "\n\nAddress where services will be provided:"
                f"\n{address(self.estimate_info.location)}"
            )
        else:
            # Takes out the initial assessment and registration fee rows when
            # the estimate is for an additional year of services.
            document.add_heading(
                "Itemized estimate for 12 session course of treatment", level=2
            )
            self.create_estimate_table(document, low_estimate)
            document.add_heading(
                "Itemized estimate for 24 session course of treatment", level=2
            )
            self.create_estimate_table(document, high_estimate)

            estimate_range_low = (int(self.calculate_low_estimate())
                    + int(self.estimate_info.session_rate) + 25)
            estimate_range_high = (int(self.calculate_high_estimate())
                + int(self.estimate_info.session_rate) + 25)
            document.add_paragraph(
                f"\nEstimate range: ${estimate_range_low}-{estimate_range_high}"
                "\n\nAddress where services will be provided:"
                f"\n{address(self.estimate_info.location)}"
            )
        runner = document.add_paragraph(
            self.create_other_sections(self.section2)
        )
        runner.add_run(
            "\n\nThis Good Faith Estimate is not a contract. It does not "
            "obligate you to accept the services listed above."
        ).bold = True

        self.record_filename()
        filepath = "/Users/RyanO/Desktop/"

        # Will need to change the save filepath for each setup environment
        document.save(
            os.path.join(filepath, self.filename)
        )

    def create_section1(self) -> str:
        """Reads text from a given file and auto-populates information."""
        lines = []
        with open(self.section1) as section1:
            for line in section1:
                if "{full_name}" in line:
                    client_full_name = (f"{self.estimate_info.client_first_name} "
                    f"{self.estimate_info.client_last_name}")
                    lines.append(
                        line.format(full_name=(client_full_name).rstrip())
                    )

                elif "{date_of_birth}" in line:
                    lines.append(
                        line.format(
                            date_of_birth=self.estimate_info.client_dob
                        ).rstrip()
                        + "\n"
                    )

                elif "{date}" in line:
                    lines.append(
                        line.format(
                            date=self.estimate_info.date_of_estimate.strftime("%x")
                            .rstrip()
                            + "\n"
                        )
                    )

                elif "{therapist_name}" in line:
                    if self.estimate_info.therapist_first_name == "Unmatched":
                        therapist_full_name = (
                            f"{self.estimate_info.therapist_first_name}\n"
                            f"EIN: N/A\n"
                        )
                    elif self.estimate_info.therapist_tax_id:
                        therapist_full_name = (
                            f"{self.estimate_info.therapist_first_name} "
                            f"{self.estimate_info.therapist_last_name}, "
                            f"{self.estimate_info.therapist_license_type}\n"
                            f"EIN: {self.estimate_info.therapist_tax_id}\n"
                        )
                    else:
                        therapist_full_name = (
                            f"{self.estimate_info.therapist_first_name} "
                            f"{self.estimate_info.therapist_last_name}, "
                            f"{self.estimate_info.therapist_license_type}\n"
                            f"EIN: N/A\n"
                        )

                    lines.append(
                        line.format(
                            therapist_name=therapist_full_name
                        ).rstrip()
                        + "\n"
                    )

                elif "{npi}" in line:
                    if self.estimate_info.therapist_npi == "":
                        therapist_npi = "N/A"
                    else:
                        therapist_npi = self.estimate_info.therapist_npi
                    lines.append(line.format(npi=therapist_npi))

                elif line == "\n":
                    lines.append("\n\n")

                else:
                    lines.append(line.rstrip() + " ")

        return "".join(lines)

    def calculate_low_estimate(self):
        return int(self.estimate_info.session_rate) * self.session_count_low

    def calculate_high_estimate(self):
        return int(self.estimate_info.session_rate) * self.session_count_high

    def create_other_sections(self, file: str) -> str:
        """Reads text from a file to use in sections of a Good Faith Estimate"""
        lines = []
        with open(file) as disclaimer:
            for line in disclaimer:
                if line == "\n":
                    lines.append(line.rstrip() + "\n\n")
                else:
                    lines.append(line.rstrip() + " ")

        return "".join(lines)

    def create_estimate_table(
        self, document_obj: "Document", itemized_list: list
    ) -> "Document.add_table":
        """Creates a table of itemized services for a Good Faith Estimate."""

        table = document_obj.add_table(rows=1, cols=6)

        headers = [
            "Service",
            "Service code",
            "Diagnosis",
            "Cost",
            "Quantity",
            "Estimate",
        ]

        header_cell = table.rows[0].cells

        for index, item in enumerate(headers):
            header_cell[index].text = item

        for i in itemized_list:
            row_cells = table.add_row().cells
            for index, item in enumerate(i):
                row_cells[index].text = item

        return table

    def record_filename(self):
        self.filename = (f"{self.estimate_info.client_last_name}_{self.estimate_info.client_first_name}"
                         f"_{self.estimate_info.date_of_estimate.strftime('%Y-%m-%d-%H-%M-%S')}.docx")


    def resource_path(self, relative_path):
        """Get the absolute path to a given resource."""

        base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))

        return os.path.join(base_path, relative_path)
